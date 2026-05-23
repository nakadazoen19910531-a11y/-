"""
ドキュメント生成サービス
ローカル環境: 既存コアエンジン (sekoplan_creator_from_notebook.py) を使用
クラウド環境: python-docx によるフォールバック生成
"""
import os
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

# ─── パス設定 ────────────────────────────────────────────────────────────────
# 環境変数 CORE_ENGINE_DIR が設定されていれば使う（ローカル開発用）
# 未設定の場合はクラウド環境としてフォールバック生成を使用
_CORE_ENGINE_DIR_STR = os.getenv('CORE_ENGINE_DIR', '')
CORE_ENGINE_DIR: Optional[Path] = Path(_CORE_ENGINE_DIR_STR) if _CORE_ENGINE_DIR_STR else None

# 出力ディレクトリ（クラウドでは /tmp 以下、ローカルでは CORE_ENGINE_DIR/output）
if CORE_ENGINE_DIR and CORE_ENGINE_DIR.exists():
    _DEFAULT_OUTPUT_DIR = CORE_ENGINE_DIR / 'output'
    sys.path.insert(0, str(CORE_ENGINE_DIR))
else:
    _DEFAULT_OUTPUT_DIR = Path(tempfile.gettempdir()) / 'sekoplan_output'


class DocumentService:
    """施工計画書 DOCX 生成サービス"""

    def __init__(self):
        self.output_dir = _DEFAULT_OUTPUT_DIR
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # テンプレートディレクトリ
        if CORE_ENGINE_DIR and CORE_ENGINE_DIR.exists():
            self.template_dir = CORE_ENGINE_DIR / 'templates' / '施工計画書'
        else:
            self.template_dir = None

        # コアエンジンの読み込み（ローカル環境のみ）
        self._core_available = False
        if CORE_ENGINE_DIR and CORE_ENGINE_DIR.exists():
            try:
                from sekoplan_creator_from_notebook import (
                    NotebookLMParser,
                    SekoplanCreatorNotebook,
                )
                self._NotebookLMParser = NotebookLMParser
                self._SekoplanCreatorNotebook = SekoplanCreatorNotebook
                self._core_available = True
                print('✅ コアエンジン読み込み成功')
            except ImportError as e:
                print(f'⚠️ コアエンジン読み込みエラー: {e}')
                print('   → フォールバックモードで起動します')
        else:
            print('ℹ️ CORE_ENGINE_DIR 未設定 → クラウドフォールバックモードで起動')

    # ─── 公開 API ──────────────────────────────────────────────────────────────
    def generate_docx(self, data: dict, template_id: Optional[str] = None) -> str:
        """DOCX を生成してファイルパスを返す。template_id が指定された場合はテンプレートを使用する"""
        # テンプレート指定がある場合は優先してテンプレートモードで生成
        if template_id:
            result = self._generate_with_template(data, template_id)
            if result:
                return result
            # テンプレート生成に失敗した場合は通常生成にフォールバック
            print(f'⚠️ テンプレート生成失敗 (id={template_id}) → フォールバックへ')

        if self._core_available:
            return self._generate_with_core_engine(data)
        return self._generate_fallback(data)

    def _generate_with_template(self, data: dict, template_id: str) -> Optional[str]:
        """アップロード済みテンプレートを使ってプレースホルダー置換で DOCX を生成する"""
        try:
            from services.template_service import TemplateService
            ts = TemplateService()

            timestamp    = datetime.now().strftime('%Y%m%d_%H%M%S')
            project_name = data.get('projectName', 'unnamed')
            safe_name    = ''.join(c for c in project_name if c not in r'\/:*?"<>|')[:50]
            filename     = f'{safe_name}_施工計画書_{timestamp}.docx'
            output_path  = self.output_dir / filename

            success = ts.apply_template(template_id, data, output_path)
            if success:
                print(f'✅ テンプレート生成成功: {output_path}')
                return str(output_path)
            return None
        except Exception as e:
            print(f'テンプレート生成エラー: {e}')
            return None

    def get_file_size(self, file_path: str) -> int:
        try:
            return os.path.getsize(file_path)
        except OSError:
            return 0

    # ─── コアエンジン生成（ローカル） ────────────────────────────────────────────
    def _generate_with_core_engine(self, data: dict) -> str:
        try:
            notebook_text = self._convert_to_notebook_format(data)
            creator = self._SekoplanCreatorNotebook()
            creator.output_dir = self.output_dir
            creator.create_from_notebook_text(notebook_text)
            output_file = self._get_latest_output_file(data.get('projectName', ''))
            if output_file:
                return str(output_file)
            return self._find_generated_file(data.get('projectName', ''))
        except Exception as e:
            print(f'コアエンジンエラー: {e} → フォールバックへ')
            return self._generate_fallback(data)

    def _convert_to_notebook_format(self, data: dict) -> str:
        """フォームデータ → NotebookLM 形式テキスト"""
        project_name    = data.get('projectName', '')
        project_type    = data.get('projectType', '')
        contract_number = data.get('contractNumber', '')
        location        = data.get('location', '')
        start_date      = data.get('startDate', '')
        end_date        = data.get('endDate', '')
        contract_amount = data.get('contractAmount', '')
        client          = data.get('client', '渋谷区')
        contractor      = data.get('contractor', '中田造園株式会社')
        work_period     = f'{start_date}から{end_date}まで'

        return f"""1. 工事件名：{project_name}

2. 工事場所：{location}

3. 契約番号：{contract_number}

4. 工事概要
工事種別: {project_type}

（受注者） 住所：東京都文京区小石川1-1-1 氏名：{contractor}
（代理人） 住所：東京都文京区小石川1-1-1 氏名：{contractor}

現場代理人・主任技術者：担当者名

工期（契約工期）：{work_period}

発注者：{client}
契約金額：{contract_amount}

5. 施工方法
施工方法の詳細を記入してください。

6. 再生資源

7. 安全管理

----------------
"""

    # ─── フォールバック生成（クラウド） ──────────────────────────────────────────
    def _generate_fallback(self, data: dict) -> str:
        """python-docx で基本的な施工計画書 DOCX を生成"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # ページ設定（A4）
            section = doc.sections[0]
            section.page_height = Cm(29.7)
            section.page_width  = Cm(21.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)

            # タイトル
            title = doc.add_heading('施 工 計 画 書', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run(data.get('projectName', ''))
            run.font.size = Pt(14)
            run.font.bold = True
            doc.add_paragraph()

            # ── 工事概要 ────────────────────────────────────────────────────────
            doc.add_heading('１　工事概要', level=1)
            rows_data = [
                ('工　事　名',   data.get('projectName',    '')),
                ('工事種別',     data.get('projectType',    '')),
                ('契約番号',     data.get('contractNumber', '')),
                ('工事場所',     data.get('location',       '')),
                ('工期（始期）', data.get('startDate',      '')),
                ('工期（終期）', data.get('endDate',        '')),
                ('契約金額',     data.get('contractAmount', '')),
                ('発注者',       data.get('client',         '')),
            ]
            tbl = doc.add_table(rows=len(rows_data), cols=2)
            tbl.style = 'Table Grid'
            for i, (label, value) in enumerate(rows_data):
                row = tbl.rows[i]
                row.cells[0].text = label
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[0].width = Cm(5)
                row.cells[1].text = value
            doc.add_paragraph()

            # ── 受注者情報 ──────────────────────────────────────────────────────
            doc.add_heading('２　受注者情報', level=1)
            contractor_rows = [
                ('受注者名', data.get('contractor', '中田造園株式会社')),
                ('住　所',   '東京都文京区小石川1-1-1'),
                ('電話番号', ''),
            ]
            tbl2 = doc.add_table(rows=len(contractor_rows), cols=2)
            tbl2.style = 'Table Grid'
            for i, (label, value) in enumerate(contractor_rows):
                row = tbl2.rows[i]
                row.cells[0].text = label
                row.cells[0].paragraphs[0].runs[0].font.bold = True
                row.cells[0].width = Cm(5)
                row.cells[1].text = value
            doc.add_paragraph()

            # ── 施工スケジュール ────────────────────────────────────────────────
            doc.add_heading('３　施工スケジュール', level=1)
            doc.add_paragraph('着工前に詳細スケジュールを作成・提出します。')
            doc.add_paragraph()

            # ── 施工方法 ────────────────────────────────────────────────────────
            doc.add_heading('４　施工方法及び留意事項', level=1)
            doc.add_paragraph('施工方法及び留意事項の詳細を記入してください。')
            doc.add_paragraph()

            # ── 安全管理 ────────────────────────────────────────────────────────
            doc.add_heading('５　安全管理計画', level=1)
            doc.add_paragraph('安全管理計画の詳細を記入してください。')
            doc.add_paragraph()

            # ── 再生資源の利用 ──────────────────────────────────────────────────
            doc.add_heading('６　再生資源の利用', level=1)
            doc.add_paragraph('再生資源の利用計画を記入してください。')

            # ファイル保存
            timestamp    = datetime.now().strftime('%Y%m%d_%H%M%S')
            project_name = data.get('projectName', 'unnamed')
            safe_name    = ''.join(c for c in project_name if c not in r'\/:*?"<>|')[:50]
            filename     = f'{safe_name}_施工計画書_{timestamp}.docx'
            file_path    = self.output_dir / filename
            doc.save(str(file_path))
            return str(file_path)

        except Exception as e:
            raise Exception(f'DOCX生成に失敗しました: {e}')

    # ─── ヘルパー ──────────────────────────────────────────────────────────────
    def _get_latest_output_file(self, project_name: str = '') -> Optional[Path]:
        try:
            docx_files = list(self.output_dir.glob('*.docx'))
            if not docx_files:
                return None
            if project_name:
                matched = [f for f in docx_files if project_name in f.name]
                if matched:
                    return max(matched, key=lambda f: f.stat().st_mtime)
            return max(docx_files, key=lambda f: f.stat().st_mtime)
        except Exception:
            return None

    def _find_generated_file(self, project_name: str) -> str:
        latest = self._get_latest_output_file(project_name)
        if latest:
            return str(latest)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return str(self.output_dir / f'{project_name}_施工計画書_{timestamp}.docx')
