"""
テンプレート管理サービス - デュアルモード実装

SUPABASE_URL / SUPABASE_KEY が設定されている場合: Supabase PostgreSQL を使用
（ファイルデータは base64 エンコードして file_data_b64 カラムに保存）
未設定の場合: JSON ファイル + ローカルファイルシステムによるフォールバック
"""
import base64
import json
import os
import uuid
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any


# JSON フォールバック用データディレクトリ
_BASE_DIR = Path(__file__).parent.parent / 'data' / 'templates'


class TemplateService:
    """DOCX テンプレートの CRUD を管理するサービス（Supabase優先 / JSONフォールバック）"""

    def __init__(self, data_dir: Optional[Path] = None):
        from db.supabase_client import get_client
        self._sb = get_client()

        if not self._sb:
            # JSON フォールバックモード
            self.data_dir = data_dir or _BASE_DIR
            self.files_dir = self.data_dir / 'files'
            self.metadata_file = self.data_dir / 'metadata.json'
            self.data_dir.mkdir(parents=True, exist_ok=True)
            self.files_dir.mkdir(parents=True, exist_ok=True)
            self._init_metadata()

    # ─── JSON フォールバック 初期化 ────────────────────────────────────────────
    def _init_metadata(self) -> None:
        if not self.metadata_file.exists():
            self.metadata_file.write_text(
                json.dumps({'templates': []}, ensure_ascii=False, indent=2),
                encoding='utf-8'
            )

    def _load_metadata(self) -> Dict[str, Any]:
        try:
            return json.loads(self.metadata_file.read_text(encoding='utf-8'))
        except Exception:
            return {'templates': []}

    def _save_metadata(self, templates: List[Dict]) -> None:
        self.metadata_file.write_text(
            json.dumps({'templates': templates}, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )

    # ─── 一覧取得 ─────────────────────────────────────────────────────────────

    def get_all(self) -> List[Dict]:
        """全テンプレートのメタデータ一覧を返す（ファイルデータは除外）"""
        if self._sb:
            try:
                res = (
                    self._sb.table('templates')
                    # file_data_b64 は大きいので除外
                    .select('id, name, description, original_filename, file_size, created_at')
                    .order('created_at', desc=True)
                    .execute()
                )
                return [self._normalize(t, file_exists=True) for t in (res.data or [])]
            except Exception as e:
                print(f'Supabase get_all テンプレートエラー: {e}')
                return []

        # JSON fallback
        templates = self._load_metadata().get('templates', [])
        result = []
        for t in templates:
            file_path = self.files_dir / f"{t['id']}.docx"
            t['file_exists'] = file_path.exists()
            result.append(t)
        return result

    # ─── 単件取得 ─────────────────────────────────────────────────────────────

    def get_by_id(self, template_id: str) -> Optional[Dict]:
        """指定 ID のテンプレートメタデータを返す（ファイルデータは除外）"""
        if self._sb:
            try:
                res = (
                    self._sb.table('templates')
                    .select('id, name, description, original_filename, file_size, created_at')
                    .eq('id', template_id)
                    .execute()
                )
                return self._normalize(res.data[0]) if res.data else None
            except Exception as e:
                print(f'Supabase get_by_id テンプレートエラー: {e}')
                return None

        # JSON fallback
        templates = self._load_metadata().get('templates', [])
        return next((t for t in templates if t['id'] == template_id), None)

    def get_file_bytes(self, template_id: str) -> Optional[bytes]:
        """テンプレートのファイルバイト列を返す（ダウンロード用）"""
        if self._sb:
            try:
                res = (
                    self._sb.table('templates')
                    .select('file_data_b64')
                    .eq('id', template_id)
                    .execute()
                )
                if not res.data or not res.data[0].get('file_data_b64'):
                    return None
                return base64.b64decode(res.data[0]['file_data_b64'])
            except Exception as e:
                print(f'Supabase get_file_bytes エラー: {e}')
                return None

        # JSON fallback
        file_path = self.files_dir / f'{template_id}.docx'
        if file_path.exists():
            return file_path.read_bytes()
        return None

    def get_file_path(self, template_id: str) -> Path:
        """
        テンプレートファイルのパスを返す（JSON フォールバック用）。
        Supabase モードでは一時ファイルに書き出したパスを返す。
        """
        if self._sb:
            file_bytes = self.get_file_bytes(template_id)
            if not file_bytes:
                return Path(f'/tmp/{template_id}.docx')  # 存在しない一時パス
            tmp = tempfile.NamedTemporaryFile(
                suffix='.docx', delete=False, prefix=f'tmpl_{template_id}_'
            )
            tmp.write(file_bytes)
            tmp.close()
            return Path(tmp.name)

        return self.files_dir / f'{template_id}.docx'

    # ─── 保存 ──────────────────────────────────────────────────────────────────

    def save(
        self,
        name: str,
        description: str,
        file_bytes: bytes,
        original_filename: str
    ) -> Dict:
        """テンプレートを保存してメタデータを返す"""
        template_id = str(uuid.uuid4())
        now = datetime.utcnow().isoformat() + 'Z'
        file_size = len(file_bytes)

        if self._sb:
            try:
                row = {
                    'id': template_id,
                    'name': name,
                    'description': description,
                    'original_filename': original_filename,
                    'file_data_b64': base64.b64encode(file_bytes).decode('ascii'),
                    'file_size': file_size,
                    'created_at': now,
                }
                self._sb.table('templates').insert(row).execute()
                return self._normalize(row, file_exists=True)
            except Exception as e:
                raise Exception(f'テンプレートの保存に失敗しました: {e}')

        # JSON fallback
        file_path = self.files_dir / f'{template_id}.docx'
        file_path.write_bytes(file_bytes)

        template: Dict = {
            'id': template_id,
            'name': name,
            'description': description,
            'original_filename': original_filename,
            'created_at': now,
            'size': file_size,
            'file_exists': True,
        }
        templates = self._load_metadata().get('templates', [])
        templates.append({k: v for k, v in template.items() if k != 'file_exists'})
        self._save_metadata(templates)
        return template

    # ─── 削除 ──────────────────────────────────────────────────────────────────

    def delete(self, template_id: str) -> bool:
        """テンプレートを削除する。成功なら True"""
        if self._sb:
            try:
                res = (
                    self._sb.table('templates')
                    .delete()
                    .eq('id', template_id)
                    .execute()
                )
                return bool(res.data)
            except Exception as e:
                print(f'Supabase delete テンプレートエラー: {e}')
                return False

        # JSON fallback
        templates = self._load_metadata().get('templates', [])
        if not any(t['id'] == template_id for t in templates):
            return False
        file_path = self.files_dir / f'{template_id}.docx'
        if file_path.exists():
            file_path.unlink()
        templates = [t for t in templates if t['id'] != template_id]
        self._save_metadata(templates)
        return True

    # ─── プレースホルダー置換 ──────────────────────────────────────────────────

    def apply_template(self, template_id: str, data: dict, output_path: Path) -> bool:
        """
        テンプレート DOCX のプレースホルダーをデータで置換して output_path に保存する。
        プレースホルダー形式: {{フィールド名}} または {{projectName}} 等
        成功なら True、失敗なら False を返す。
        """
        try:
            from docx import Document
        except ImportError:
            return False

        # ファイルパスを取得（Supabase モードでは一時ファイル）
        file_path = self.get_file_path(template_id)
        _is_temp = self._sb is not None  # Supabase モードでは一時ファイルを後で削除

        if not file_path.exists():
            return False

        try:
            doc = Document(str(file_path))
        except Exception:
            if _is_temp and file_path.exists():
                file_path.unlink()
            return False

        replacements = self._build_replacements(data)

        # 段落のプレースホルダーを置換
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # テーブル内のプレースホルダーを置換
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        # ヘッダー・フッター内も置換
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if header_footer:
                    for paragraph in header_footer.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        doc.save(str(output_path))

        # Supabase モードの一時ファイルを削除
        if _is_temp and file_path.exists():
            try:
                file_path.unlink()
            except Exception:
                pass

        return True

    # ─── ヘルパー ─────────────────────────────────────────────────────────────

    def _normalize(self, row: dict, file_exists: bool = True) -> Dict:
        """Supabase 行データをフロントエンド向けフォーマットに変換"""
        return {
            'id': row.get('id', ''),
            'name': row.get('name', ''),
            'description': row.get('description', ''),
            'original_filename': row.get('original_filename', ''),
            'size': row.get('file_size', row.get('size', 0)),
            'created_at': row.get('created_at', ''),
            'file_exists': file_exists,
        }

    def _build_replacements(self, data: dict) -> Dict[str, str]:
        """フォームデータからプレースホルダー置換マッピングを構築する"""
        field_map = {
            'projectName':    ['projectName', '工事名', '工事件名'],
            'projectType':    ['projectType', '工事種別'],
            'contractNumber': ['contractNumber', '契約番号'],
            'location':       ['location', '工事場所'],
            'startDate':      ['startDate', '工期始期', '着工日'],
            'endDate':        ['endDate', '工期終期', '竣工日'],
            'contractAmount': ['contractAmount', '契約金額'],
            'client':         ['client', '発注者'],
            'contractor':     ['contractor', '受注者'],
        }
        replacements: Dict[str, str] = {}
        for field_key, placeholder_keys in field_map.items():
            value = data.get(field_key, '')
            for pk in placeholder_keys:
                replacements[f'{{{{{pk}}}}}'] = value  # {{key}}
        return replacements

    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> None:
        """段落内のプレースホルダーをテキスト置換する（run をまたぐ場合にも対応）"""
        full_text = ''.join(run.text for run in paragraph.runs)
        new_text = full_text
        for placeholder, value in replacements.items():
            new_text = new_text.replace(placeholder, value)

        if new_text == full_text:
            return  # 変更なし

        # run[0] に全テキストを入れ、残りをクリア
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''
